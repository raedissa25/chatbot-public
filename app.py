import os
from datetime import datetime
import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import google.generativeai as genai
from dotenv import load_dotenv
from tensorflow.keras.models import load_model
import numpy as np
import streamlit as st
api_key = st.secrets["GEMINI_API_KEY"]
# Charger le modèle CNN entraîné
cnn_model = load_model("models/resnet50v2_ecg_best_model.h5")

# Classes (à adapter selon ton dataset)
class_labels = ['AHB', 'HMI', 'MI', 'Normal']

# Charger les variables d'environnement
api_key = st.secrets["GEMINI_API_KEY"]
if api_key is None:
    raise ValueError("GEMINI_API_KEY is not set in environment variables")
genai.configure(api_key=api_key)

# Paramètres du modèle
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
)
def predict_ecg_class(image_file):
    image = Image.open(image_file).convert('RGB').resize((224, 224))  
    image_array = np.array(image) / 255.0
    image_array = np.expand_dims(image_array, axis=0)  
    prediction = cnn_model.predict(image_array)
    predicted_class = class_labels[np.argmax(prediction)]
    return predicted_class

# Fonction pour analyser une image ECG et générer un rapport
def generate_ecg_details(ecg_image):
    image = Image.open(ecg_image)
    current_date = datetime.now().strftime('%Y-%m-%d')

    prompt = f"""Analysez cette image ECG et fournissez un rapport détaillé. Suivez cette structure :

**RAPPORT D'ANALYSE ECG**

**1. Informations du patient:**
- Nom:
- Age:
- Sexe:
- Numéro d'identification:
- Date de l'ECG:

**2. Informations cliniques:**
- Motif de l'ECG:
- Antécédents médicaux pertinents:
- Médicaments:

**3. DÉTAILS TECHNIQUES DE L’ECG:**
- Appareil ECG utilisé:
- Configuration des dérivations:
- Calibration:
- Qualité d'enregistrement:

**4. Résultats de l'ECG:**
**Rythme et fréquence:**
- Fréquence cardiaque:
- Rythme:
- Ondes P:
- Intervalle PR:
- Complexe QRS:
- Intervalle QT/QTc:
- Segment ST:
- Ondes T:

**Axes:**
- Axe des ondes P:
- Axe du complexe QRS:
- Axe des ondes T:

**Conduction et morphologie:**
- Conduction atriale:
- Conduction ventriculaire:
- Morphologie du complexe QRS:
- Modifications du segment ST-T:

**5. INTERPRETATION:**
- Normal ou anormal:
- Diagnostic / Résultats:
- Comparaison avec l’ECG précédent (si disponible):

**6. CONCLUSION ET RECOMMANDATIONS:**
- Résumé:
- Recommandations:

**7. CARDIOLOGUE RÉDACTEUR DU RAPPORT:**
- Nom:
- Signature: Impossible de fournir une signature pour un rapport généré par IA.
- Date of Report: {current_date}
"""

    chat_session = model.start_chat(history=[])
    predicted_class = predict_ecg_class(ecg_image)
    full_prompt = prompt + f"\n\n**Classe d’anomalie (predite par notre modele Resnet50v2):** {predicted_class}\n\nNow complete the rest of the report using the above prediction as reference."
    response = chat_session.send_message([full_prompt, image])
    return response.text

# Fonction pour créer un document Word contenant le rapport ECG
def create_doc(report_text, ecg_image):
    doc = Document()
    doc.add_heading('RAPPORT D’ANALYSE DE L’ECG', 0)

    for line in report_text.split("\n"):
        if line.strip() == '':
            continue
        if line.startswith('**') and line.endswith('**'):
            doc.add_heading(line.strip('**'), level=1)
        elif line.startswith('-'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())

    doc.add_heading('ECG Tracing:', level=1)
    image_stream = BytesIO(ecg_image.getvalue())
    doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Interface utilisateur avec Streamlit
def main():
    st.title("🫀Chatbot de santé cardiaque – Obtenez une analyse ECG instantanée")

    # Section Upload ECG
    st.header("📂Téléverser l’image de l’ECG")
    ecg_image = st.file_uploader("Téléversez une image de l’ECG", type=["png", "jpg", "jpeg"])

    if ecg_image is not None:
        st.image(ecg_image, caption='✅Image ECG téléversée', use_column_width=True)

        if st.button("Génerer un rapport ECG"):
            with st.spinner("🔍Analyse de l’image ECG..."):
                ecg_details = generate_ecg_details(ecg_image)
            st.header("📊Rapport ECG généré")
            st.markdown(ecg_details)

            # Stocker le rapport dans la session pour le téléchargement
            st.session_state.ecg_details = ecg_details

        # Bouton de téléchargement du rapport
        if hasattr(st.session_state, 'ecg_details'):
            doc_file_stream = create_doc(st.session_state.ecg_details, ecg_image)
            st.download_button(
                label="📥Télécharger le rapport ECG",
                data=doc_file_stream,
                file_name="ECG_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Section Chatbot IA
    st.header("💬Posez votre question à votre cardiologue IA ")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Afficher l'historique du chat
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Saisie utilisateur
    user_input = st.chat_input("Posez votre question...")
    if user_input:
        st.session_state.messages.append({"role": "user", "content": user_input})
        
        with st.chat_message("user"):
            st.markdown(user_input)

        with st.spinner("En réflexion..."):
            chat_session = model.start_chat(history=[])
            response = chat_session.send_message(user_input)
            bot_response = response.text

        st.session_state.messages.append({"role": "assistant", "content": bot_response})

        with st.chat_message("assistant"):
            st.markdown(bot_response)

if __name__ == '__main__':
    main()
